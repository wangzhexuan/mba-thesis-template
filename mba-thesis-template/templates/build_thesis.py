#!/usr/bin/env python3
"""
MBA论文Word文档生成脚本

功能：
- 根据Markdown草稿生成规范的Word文档
- 应用学校格式规范
- 生成目录、页眉页脚
- 插入图表和参考文献

使用方法：
    python build_thesis.py

依赖：
    pip install python-docx pyyaml markdown
"""

import os
import re
import yaml
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("请安装python-docx: pip install python-docx")
    exit(1)


class ThesisBuilder:
    """MBA论文构建器"""

    def __init__(self, config_dir="config", drafts_dir="drafts", output_dir="output"):
        self.config_dir = Path(config_dir)
        self.drafts_dir = Path(drafts_dir)
        self.output_dir = Path(output_dir)

        # 加载配置
        self.thesis_config = self.load_config("thesis-config.yaml")
        self.format_config = self.load_config("school-format.yaml")

        # 初始化文档
        self.doc = Document()

    def load_config(self, filename):
        """加载YAML配置文件"""
        config_path = self.config_dir / filename
        if config_path.exists():
            with open(config_path, 'r', encoding='utf-8') as f:
                return yaml.safe_load(f)
        return {}

    def setup_document(self):
        """设置文档基本格式"""
        # 设置页面边距
        sections = self.doc.sections
        for section in sections:
            margin = self.format_config.get('format', {}).get('margin', {})
            section.top_margin = Cm(margin.get('top', 2.54))
            section.bottom_margin = Cm(margin.get('bottom', 2.54))
            section.left_margin = Cm(margin.get('left', 3.17))
            section.right_margin = Cm(margin.get('right', 3.17))

        # 设置默认字体
        style = self.doc.styles['Normal']
        font = style.font
        font_config = self.format_config.get('format', {}).get('font', {})
        font.name = font_config.get('english', 'Times New Roman')
        font.size = Pt(self.format_config.get('format', {}).get('size', {}).get('body', 12))
        style.element.rPr.rFonts.set(qn('w:eastAsia'), font_config.get('body', '宋体'))

        # 设置行距
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = self.format_config.get('format', {}).get('spacing', {}).get('line', 1.5)

    def create_styles(self):
        """创建自定义样式"""
        # 一级标题
        style = self.doc.styles.add_style('Thesis Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = self.format_config.get('format', {}).get('font', {}).get('english', 'Times New Roman')
        font.size = Pt(self.format_config.get('format', {}).get('size', {}).get('title', 16))
        font.bold = True
        style.element.rPr.rFonts.set(qn('w:eastAsia'), self.format_config.get('format', {}).get('font', {}).get('title', '黑体'))
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

        # 二级标题
        style = self.doc.styles.add_style('Thesis Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = self.format_config.get('format', {}).get('font', {}).get('english', 'Times New Roman')
        font.size = Pt(self.format_config.get('format', {}).get('size', {}).get('subtitle', 14))
        font.bold = True
        style.element.rPr.rFonts.set(qn('w:eastAsia'), self.format_config.get('format', {}).get('font', {}).get('subtitle', '黑体'))
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

        # 三级标题
        style = self.doc.styles.add_style('Thesis Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = self.format_config.get('format', {}).get('font', {}).get('english', 'Times New Roman')
        font.size = Pt(self.format_config.get('format', {}).get('size', {}).get('sub_subtitle', 12))
        font.bold = True
        style.element.rPr.rFonts.set(qn('w:eastAsia'), self.format_config.get('format', {}).get('font', {}).get('sub_subtitle', '黑体'))
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)

    def add_cover(self):
        """添加封面"""
        metadata = self.thesis_config.get('metadata', {})

        # 空行
        for _ in range(3):
            self.doc.add_paragraph()

        # 学校名称
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(metadata.get('school', ''))
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 论文类型
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('工商管理硕士（MBA）学位论文')
        run.font.size = Pt(18)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 空行
        for _ in range(2):
            self.doc.add_paragraph()

        # 论文题目
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(metadata.get('title', ''))
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 空行
        for _ in range(3):
            self.doc.add_paragraph()

        # 作者信息
        info_items = [
            f"作者姓名：{metadata.get('author', '')}",
            f"学号：{metadata.get('student_id', '')}",
            f"学院：{metadata.get('college', '')}",
            f"导师：{metadata.get('supervisor', '')}",
            f"完成日期：{metadata.get('completion_date', '')}",
        ]

        for item in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(item)
            run.font.size = Pt(14)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 分页
        self.doc.add_page_break()

    def add_abstract(self):
        """添加摘要"""
        # 中文摘要
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('摘  要')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 摘要内容（从文件读取或占位）
        abstract_path = self.drafts_dir / "abstract.md"
        if abstract_path.exists():
            with open(abstract_path, 'r', encoding='utf-8') as f:
                content = f.read()
                p = self.doc.add_paragraph(content)
                p.paragraph_format.first_line_indent = Cm(0.74)
        else:
            p = self.doc.add_paragraph('[请在此处填写中文摘要]')
            p.paragraph_format.first_line_indent = Cm(0.74)

        # 关键词
        keywords = self.thesis_config.get('keywords', {}).get('chinese', [])
        p = self.doc.add_paragraph()
        run = p.add_run('关键词：')
        run.font.bold = True
        run = p.add_run('；'.join(keywords) if keywords else '[关键词]')

        # 分页
        self.doc.add_page_break()

        # 英文摘要
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Abstract')
        run.font.size = Pt(16)
        run.font.bold = True

        # 英文摘要内容
        abstract_en_path = self.drafts_dir / "abstract_en.md"
        if abstract_en_path.exists():
            with open(abstract_en_path, 'r', encoding='utf-8') as f:
                content = f.read()
                self.doc.add_paragraph(content)
        else:
            self.doc.add_paragraph('[Please write the English abstract here]')

        # 英文关键词
        keywords_en = self.thesis_config.get('keywords', {}).get('english', [])
        p = self.doc.add_paragraph()
        run = p.add_run('Keywords: ')
        run.font.bold = True
        run = p.add_run('; '.join(keywords_en) if keywords_en else '[Keywords]')

        # 分页
        self.doc.add_page_break()

    def add_chapter(self, chapter_num, chapter_title):
        """添加章节"""
        # 章标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'第{chapter_num}章 {chapter_title}')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.style = self.doc.styles['Thesis Heading 1']

    def add_section(self, section_num, section_title, level=2):
        """添加节"""
        if level == 2:
            style_name = 'Thesis Heading 2'
        elif level == 3:
            style_name = 'Thesis Heading 3'
        else:
            style_name = 'Thesis Heading 2'

        p = self.doc.add_paragraph(f'{section_num} {section_title}')
        p.style = self.doc.styles[style_name]

    def add_paragraph(self, text, first_indent=True):
        """添加段落"""
        p = self.doc.add_paragraph(text)
        if first_indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        return p

    def add_image(self, image_path, caption=None, width=None):
        """添加图片"""
        if os.path.exists(image_path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if width:
                run = p.add_run()
                run.add_picture(image_path, width=Inches(width))
            else:
                run = p.add_run()
                run.add_picture(image_path, width=Inches(5))

            if caption:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(caption)
                run.font.size = Pt(10)

    def add_table(self, data, caption=None):
        """添加表格"""
        if caption:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.font.size = Pt(10)
            run.font.bold = True

        if data and len(data) > 0:
            table = self.doc.add_table(rows=len(data), cols=len(data[0]))
            table.style = 'Table Grid'

            for i, row in enumerate(data):
                for j, cell_text in enumerate(row):
                    cell = table.cell(i, j)
                    cell.text = str(cell_text)
                    # 设置单元格字体
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

    def add_toc(self):
        """添加目录"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('目  录')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 注意：python-docx无法自动生成目录
        # 需要在Word中手动更新目录
        self.doc.add_paragraph('[请在Word中插入自动目录：引用 -> 目录 -> 自动目录]')

        # 分页
        self.doc.add_page_break()

    def add_references(self):
        """添加参考文献"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('参考文献')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 从文件读取参考文献
        refs_path = self.output_dir / "references.bib"
        if refs_path.exists():
            # 这里简化处理，实际应该解析BibTeX
            self.doc.add_paragraph('[参考文献列表将从references.bib生成]')
        else:
            self.doc.add_paragraph('[请在此处添加参考文献]')

    def load_chapter_content(self, chapter_file):
        """加载章节内容"""
        chapter_path = self.drafts_dir / chapter_file
        if chapter_path.exists():
            with open(chapter_path, 'r', encoding='utf-8') as f:
                return f.read()
        return None

    def parse_markdown(self, content):
        """简单解析Markdown内容"""
        lines = content.split('\n')
        current_text = []

        for line in lines:
            line = line.strip()
            if not line:
                if current_text:
                    self.add_paragraph(' '.join(current_text))
                    current_text = []
                continue

            if line.startswith('# '):
                if current_text:
                    self.add_paragraph(' '.join(current_text))
                    current_text = []
                # 一级标题已在add_chapter中处理
            elif line.startswith('## '):
                if current_text:
                    self.add_paragraph(' '.join(current_text))
                    current_text = []
                title = line[3:]
                self.add_section('', title, level=2)
            elif line.startswith('### '):
                if current_text:
                    self.add_paragraph(' '.join(current_text))
                    current_text = []
                title = line[4:]
                self.add_section('', title, level=3)
            elif line.startswith('- ') or line.startswith('* '):
                # 列表项
                current_text.append(line[2:])
            else:
                current_text.append(line)

        if current_text:
            self.add_paragraph(' '.join(current_text))

    def build(self):
        """构建完整论文"""
        print("开始构建MBA论文...")

        # 设置文档
        self.setup_document()
        self.create_styles()

        # 添加封面
        print("  添加封面...")
        self.add_cover()

        # 添加摘要
        print("  添加摘要...")
        self.add_abstract()

        # 添加目录
        print("  添加目录...")
        self.add_toc()

        # 添加各章节
        chapters = self.thesis_config.get('structure', {}).get('chapters', [
            {'number': 1, 'title': '绪论'},
            {'number': 2, 'title': '相关理论与文献综述'},
            {'number': 3, 'title': '现状分析'},
            {'number': 4, 'title': '策略优化'},
            {'number': 5, 'title': '结论与展望'},
        ])

        chapter_files = [
            'chapter1-introduction.md',
            'chapter2-literature.md',
            'chapter3-analysis.md',
            'chapter4-strategy.md',
            'chapter5-conclusion.md',
        ]

        for i, chapter in enumerate(chapters):
            print(f"  添加第{chapter['number']}章：{chapter['title']}...")
            self.add_chapter(chapter['number'], chapter['title'])

            # 加载章节内容
            if i < len(chapter_files):
                content = self.load_chapter_content(chapter_files[i])
                if content:
                    self.parse_markdown(content)
                else:
                    self.add_paragraph(f'[请在 {chapter_files[i]} 中编写第{chapter["number"]}章内容]')

            # 分页
            if i < len(chapters) - 1:
                self.doc.add_page_break()

        # 添加参考文献
        print("  添加参考文献...")
        self.doc.add_page_break()
        self.add_references()

        # 保存文档
        output_path = self.output_dir / "thesis.docx"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"\n论文已生成：{output_path}")
        print("\n后续步骤：")
        print("1. 在Word中打开thesis.docx")
        print("2. 更新目录（引用 -> 目录 -> 更新目录）")
        print("3. 检查格式是否符合学校要求")
        print("4. 添加图表和公式")


def main():
    """主函数"""
    builder = ThesisBuilder()
    builder.build()


if __name__ == "__main__":
    main()
